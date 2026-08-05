# -*- coding: utf-8 -*-
"""Build Literature_Review_BUSI.docx (text only) from curated content."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROJ = os.path.dirname(ROOT)
OUT  = os.path.join(PROJ, "Literature_Review_BUSI.docx")

doc = Document()
# base style
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8); normal.paragraph_format.line_spacing = 1.15

def H1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1F,0x3A,0x5F)
    return h
def H2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.color.rgb = RGBColor(0x2E,0x50,0x77)
    return h
def P(t, justify=True):
    p = doc.add_paragraph(t)
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

# ---------------- Title ----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("A Literature Review of Deep Learning for Breast Cancer\nClassification and Segmentation on Ultrasound Images")
r.bold = True; r.font.size = Pt(16); r.font.name = "Times New Roman"
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("COMP9444 Group Project 047 — Breast Ultrasound (BUSI) Dataset")
rs.italic = True; rs.font.size = Pt(11)
doc.add_paragraph()

# ---------------- 1. Introduction ----------------
H1("1. Introduction")
P("Breast cancer is among the leading causes of cancer-related death in women worldwide, and early "
  "detection remains the single most effective lever for reducing mortality and improving prognosis. "
  "Ultrasound imaging is a safe, inexpensive, radiation-free and widely available modality for breast "
  "examination, and it is particularly valuable for dense breasts where mammography is less sensitive. "
  "However, ultrasound images are notoriously difficult to interpret: speckle noise, low contrast, "
  "acoustic shadowing, and the large variability of lesion appearance make manual reading operator-"
  "dependent and time-consuming, even for experienced radiologists [18][19]. These challenges have "
  "motivated a large body of work on computer-aided diagnosis (CAD), in which machine learning, and "
  "more recently deep learning, is used to automatically classify lesions (normal, benign, malignant) "
  "and to segment lesion boundaries.")
P("This review surveys the literature that underpins Project 047, whose goal is to build a deep-learning "
  "pipeline for breast cancer classification and segmentation on the Breast Ultrasound Images (BUSI) "
  "dataset [1]. We organise the discussion around six themes: (i) the dataset and its data-quality "
  "considerations; (ii) convolutional backbones for image classification; (iii) semantic segmentation "
  "architectures from natural and medical imaging; (iv) domain-specific breast-ultrasound segmentation "
  "and classification methods; (v) unified two-stage and multi-task pipelines; and (vi) strong general "
  "baselines, foundation models, and explainability. We conclude by identifying the research gap that "
  "this project addresses and by motivating our methodological choices.")

# ---------------- 2. Dataset ----------------
H1("2. The BUSI Dataset and Data Resources")
P("The Breast Ultrasound Images (BUSI) dataset of Al-Dhabyani et al. [1] is the empirical foundation of "
  "this project. It contains 780 ultrasound images collected in 2018 from 600 female patients aged "
  "25–75 at Baheya Hospital (Cairo, Egypt), categorised into three classes — normal, benign and "
  "malignant — with an average size of roughly 500x500 pixels and stored in PNG format. Crucially, "
  "BUSI provides pixel-level ground-truth lesion masks in addition to image-level class labels, which is "
  "what makes it suitable for both classification and segmentation. This dual annotation has made BUSI "
  "one of the most widely used public benchmarks in the field.")
P("The same group earlier explored data augmentation and classification of breast masses using CNNs on "
  "ultrasound images [24], demonstrating that conventional and GAN-based augmentation can partially "
  "compensate for the small size of ultrasound datasets. Their findings directly justify the augmentation "
  "strategy adopted in this project and highlight the recurring small-sample problem in the domain.")
P("Because BUSI is small and heterogeneous, careful data curation is essential. Our own exploratory data "
  "analysis, and independent reports in the literature, reveal that the dataset is class-imbalanced "
  "(benign > malignant > normal), that images are not truly uniform in size, that some images carry "
  "burned-in annotations and calipers, and — most importantly — that BUSI contains exact and near-"
  "duplicate images, including a small number of duplicates that appear under conflicting class labels. "
  "These issues can leak information between training and test partitions and inflate reported "
  "performance, and they must be handled through de-duplication and grouped, leakage-free data splitting. "
  "This concern is consistent with community discussion of the dataset's limitations and with the broader "
  "observation that unusually high accuracies reported on small ultrasound datasets should be treated "
  "with caution [14].")

# ---------------- 3. Classification backbones ----------------
H1("3. Convolutional Backbones for Image Classification")
P("Modern image classification is dominated by deep convolutional neural networks (CNNs), and three "
  "backbone families named in the project brief — ResNet, DenseNet and EfficientNet — form the basis of "
  "our classification experiments.")
H2("3.1 Residual and densely connected networks")
P("He et al. [4] introduced residual learning (ResNet), in which identity shortcut connections allow "
  "gradients to flow through very deep networks and mitigate the degradation problem. ResNet won the "
  "ILSVRC-2015 classification challenge and remains a strong, general-purpose backbone. Huang et al. [5] "
  "extended this idea with DenseNet, connecting each layer to every subsequent layer so that features are "
  "reused throughout the network; this yields competitive accuracy with fewer parameters, an attractive "
  "property when training data are scarce, as is the case for breast ultrasound.")
H2("3.2 Compound scaling with EfficientNet")
P("Tan and Le [6] proposed EfficientNet, which uses a principled compound-scaling rule to jointly scale "
  "network depth, width and input resolution. EfficientNet models achieve state-of-the-art accuracy with "
  "substantially fewer parameters and FLOPs than earlier architectures. For small medical datasets the "
  "lower-capacity variants (e.g. EfficientNet-B0) are especially appealing because they are less prone to "
  "over-fitting while still benefiting from ImageNet pre-training.")
H2("3.3 Application to breast ultrasound classification")
P("These backbones have been transferred directly to breast ultrasound. Latha et al. [14] fine-tuned "
  "EfficientNet-B7 for three-class BUSI classification, combined it with preprocessing (Gaussian and "
  "Sobel filtering, histogram equalisation), augmentation and Grad-CAM explainability, and reported very "
  "high overall accuracy (~99%). While encouraging, such near-perfect scores on a small dataset should be "
  "interpreted cautiously in light of the duplication and leakage issues noted above; they underline the "
  "importance of rigorous, leakage-free evaluation rather than headline accuracy alone.")

# ---------------- 4. Segmentation architectures ----------------
H1("4. Semantic Segmentation Architectures")
P("Lesion segmentation delineates the pixel-level boundary of a mass and is central to this project. We "
  "review the canonical encoder-decoder families and their medical-imaging refinements, together with the "
  "atrous-convolution and Transformer lines of work named in the brief.")
H2("4.1 U-Net and its variants")
P("The U-Net of Ronneberger et al. [3] is the archetypal biomedical segmentation network: a contracting "
  "encoder captures context, an expanding decoder restores spatial resolution, and skip connections carry "
  "fine-grained detail across the two paths, enabling accurate segmentation from relatively few annotated "
  "images. U-Net is our first segmentation baseline. Several refinements target its known weaknesses. "
  "UNet++ [8] introduces nested, densely connected skip pathways and deep supervision to reduce the "
  "semantic gap between encoder and decoder features. UNet 3+ [10] goes further with full-scale skip "
  "connections that fuse low-level detail and high-level semantics across all resolutions, which is "
  "helpful when target sizes vary widely — as breast lesions do. Attention U-Net [9] augments the skip "
  "connections with attention gates that learn to focus on salient target regions and suppress "
  "irrelevant background; because breast-ultrasound lesions are often small and embedded in noisy "
  "tissue, this attention mechanism is particularly well suited to our task and is our primary "
  "segmentation improvement over the plain U-Net baseline.")
H2("4.2 Atrous convolution: the DeepLab family")
P("A complementary line of work replaces the U-Net decoder with atrous (dilated) convolutions. Chen et "
  "al. [7] proposed DeepLabv3+, which combines atrous spatial pyramid pooling (ASPP) for multi-scale "
  "context with a lightweight decoder that sharpens object boundaries, and uses atrous separable "
  "convolutions for efficiency. DeepLabv3+ achieves strong results on natural-image benchmarks and, in "
  "this project, serves as a second segmentation baseline alongside U-Net, as suggested by the brief.")
H2("4.3 Transformer-based segmentation")
P("More recently, Transformers have been introduced to capture long-range dependencies that convolutions "
  "model only locally. TransUNet [11] couples a CNN feature extractor with a Transformer encoder and a "
  "U-Net-style decoder, reporting strong medical-segmentation performance. Transformers, however, are "
  "data-hungry and typically require large-scale pre-training; on a small dataset such as BUSI they are "
  "prone to over-fitting. We therefore treat TransUNet as a point of comparison and future direction "
  "rather than a primary model.")

# ---------------- 5. Domain-specific ----------------
H1("5. Domain-Specific Breast-Ultrasound Methods")
P("Beyond generic architectures, a growing literature adapts deep learning specifically to breast "
  "ultrasound. Yap et al. [2] were among the first to apply CNNs (a patch-based LeNet, a U-Net and a "
  "transfer-learned FCN-AlexNet) to breast-ultrasound lesion detection, establishing that deep models "
  "outperform classical hand-crafted pipelines for locating lesions. Subsequent work has focused on "
  "handling the modality's characteristic difficulties — blurred boundaries, low signal-to-noise ratio, "
  "intensity heterogeneity and irregular malignant shapes.")
P("Chen et al. [12] proposed AAU-Net, which replaces ordinary convolutions with a hybrid adaptive "
  "attention module that integrates channel and spatial self-attention; it was designed explicitly for "
  "breast-ultrasound lesion segmentation and reported improved accuracy and generalisation over several "
  "state-of-the-art methods on public datasets. Derakhshandeh and Mahloojifar [13] introduced CResU-Net, "
  "a U-Net-based encoder-decoder that combines residual and MultiResUNet ideas with a 'Co-Block', and "
  "reported a Dice coefficient of about 82.9% on BUSI while remaining relatively lightweight. Sulaiman "
  "et al. [15] proposed an attention-driven U-Net evaluated directly on BUSI (with the dataset balanced "
  "through augmentation), confirming the benefit of attention for this task, although some reported "
  "figures vary between tables and should be quoted carefully. Anari et al. [16] combined a modified "
  "U-Net with ResNet-18, DenseNet-121 and EfficientNet-B0 encoders together with CBAM and non-local "
  "attention and Grad-CAM explanations; their configuration is closely aligned with the backbone set of "
  "Project 047, though the reported Dice (~0.61) is more modest, illustrating that BUSI segmentation "
  "remains an open problem and that very high scores are not universal.")

# ---------------- 6. Unified pipelines ----------------
H1("6. Unified Two-Stage and Multi-Task Pipelines")
P("Because Project 047 must perform both classification and segmentation, methods that unify the two "
  "tasks are of particular interest. Bruno et al. [17] proposed a dual-stage framework that first "
  "segments the suspicious region (using DeepLabv3+ with a ResNet-34 encoder), extracts the lesion "
  "region-of-interest from the predicted mask, and then classifies the cropped ROI (using compact "
  "backbones such as MobileNetV3-Small and EfficientNet-B0). Trained with a patient-level split, it "
  "reported a segmentation Dice of about 0.77 and a classification AUC of about 0.99 for the "
  "benign-versus-malignant task. This 'segment-then-classify' design is attractive for our project "
  "because its components are exactly the backbones we already implement, effectively assembling a "
  "state-of-the-art pipeline from familiar building blocks; we adopt it as the template for our main "
  "method rather than a mere baseline.")
P("An alternative route to integration is multi-task learning. Lu et al. [25] performed joint "
  "segmentation and classification of breast ultrasound images with a shared encoder feeding both a "
  "segmentation head and a classification head, coupled with object-contextual attention, and reported "
  "that jointly learning the two tasks improves both relative to single-task baselines. Multi-task "
  "learning with a shared encoder is therefore a compelling second design option, and we discuss it "
  "alongside the two-stage pipeline when motivating our final architecture.")

# ---------------- 7. Strong baselines, foundation models, XAI ----------------
H1("7. Strong Baselines, Foundation Models and Explainability")
P("Two further strands of work frame the current state of the art. First, nnU-Net [20] is a "
  "self-configuring segmentation framework that automatically adapts preprocessing, network topology, "
  "training and post-processing to a given dataset, and has surpassed many specialised solutions across "
  "23 public biomedical-segmentation challenges without manual tuning. It is widely regarded as the "
  "de-facto strong baseline in medical segmentation; although its automated pipeline is "
  "computationally heavy for a small project, it provides a rigorous performance reference.")
P("Second, promptable foundation models have begun to reshape segmentation. The Segment Anything Model "
  "(SAM) [21] is a promptable model trained on a billion masks with strong zero-shot transfer, and "
  "MedSAM [22] adapts this recipe to medicine by fine-tuning on more than a million medical image-mask "
  "pairs across many modalities. These models generalise impressively but are not optimised for breast "
  "ultrasound out of the box and typically require prompts or fine-tuning; we therefore regard them as a "
  "promising future direction, potentially as a zero-shot demonstration rather than a trained baseline.")
P("Finally, interpretability is essential for clinical acceptance. Grad-CAM [23] produces "
  "class-discriminative localisation heatmaps for any CNN without retraining and is the explainability "
  "method underlying several breast-ultrasound studies [14][16]. We incorporate Grad-CAM to visualise "
  "which image regions drive our classifier's decisions, adding a layer of transparency to the pipeline.")

# ---------------- 8. Surveys ----------------
H1("8. Surveys and Open Challenges")
P("Two surveys situate the above methods historically. Cheng et al. [19] reviewed pre-deep-learning CAD "
  "pipelines for breast ultrasound — preprocessing, segmentation, feature extraction and classification "
  "— establishing the classical baseline against which deep learning is measured. Xian et al. [18] "
  "surveyed roughly two decades of breast-ultrasound segmentation, categorising methods and identifying "
  "the field's recurring challenges: speckle noise, low contrast and fuzzy boundaries. Together they "
  "explain why breast-ultrasound analysis is intrinsically difficult and why modern, context- and "
  "attention-aware deep models are needed.")

# ---------------- 9. Synthesis ----------------
H1("9. Synthesis and Research Gap")
P("The literature yields a clear picture. Generic backbones (ResNet [4], DenseNet [5], EfficientNet [6]) "
  "and segmentation architectures (U-Net [3], its variants [8][9][10], and DeepLabv3+ [7]) provide "
  "well-understood, transferable building blocks. Domain-specific studies [2][12][13][15][16] confirm "
  "that attention mechanisms and boundary-aware designs help with the modality's small, blurry lesions, "
  "but they also report a wide and sometimes inconsistent range of results, indicating that BUSI "
  "segmentation is far from solved. Integrated pipelines [17][25] point toward combining segmentation and "
  "classification, while foundation models [21][22] and strong baselines [20] mark the outer edge of the "
  "field. Across all of this work, a recurring methodological weakness is insufficient attention to data "
  "quality: small datasets, duplicates, and non-grouped splits can produce optimistic results [14].")
P("This project therefore pursues two complementary aims. Methodologically, we establish faithful, "
  "leakage-free baselines: three classifiers (ResNet, DenseNet, EfficientNet-B0) and two-to-three "
  "segmentation models (U-Net and DeepLabv3+, plus Attention U-Net as a domain-motivated improvement), "
  "all evaluated under a grouped, de-duplicated data split with per-class metrics. Architecturally, we "
  "assemble these components into a segmentation-guided classification pipeline in the spirit of Bruno "
  "et al. [17], with multi-task learning [25] as an alternative design, and we add Grad-CAM [23] for "
  "interpretability. In doing so, the project contributes not only competitive models but also a "
  "rigorous, reproducible evaluation of breast-ultrasound classification and segmentation on a dataset "
  "whose limitations are explicitly acknowledged and controlled for.")

# ---------------- References ----------------
doc.add_page_break()
H1("References")
refs = [
 "[1] Al-Dhabyani, W., Gomaa, M., Khaled, H., & Fahmy, A. (2020). Dataset of breast ultrasound images. Data in Brief, 28, 104863. https://doi.org/10.1016/j.dib.2019.104863",
 "[2] Yap, M. H., Pons, G., Marti, J., Ganau, S., Sentis, M., Zwiggelaar, R., Davison, A. K., & Marti, R. (2018). Automated breast ultrasound lesions detection using convolutional neural networks. IEEE Journal of Biomedical and Health Informatics, 22(4), 1218-1226. https://doi.org/10.1109/JBHI.2017.2731873",
 "[3] Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional networks for biomedical image segmentation. In MICCAI (pp. 234-241). Springer. https://doi.org/10.1007/978-3-319-24574-4_28",
 "[4] He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image recognition. In CVPR (pp. 770-778). https://doi.org/10.1109/CVPR.2016.90",
 "[5] Huang, G., Liu, Z., van der Maaten, L., & Weinberger, K. Q. (2017). Densely connected convolutional networks. In CVPR (pp. 4700-4708). https://doi.org/10.1109/CVPR.2017.243",
 "[6] Tan, M., & Le, Q. V. (2019). EfficientNet: Rethinking model scaling for convolutional neural networks. In ICML (pp. 6105-6114). arXiv:1905.11946",
 "[7] Chen, L.-C., Zhu, Y., Papandreou, G., Schroff, F., & Adam, H. (2018). Encoder-decoder with atrous separable convolution for semantic image segmentation (DeepLabv3+). In ECCV (pp. 801-818). arXiv:1802.02611",
 "[8] Zhou, Z., Rahman Siddiquee, M. M., Tajbakhsh, N., & Liang, J. (2018). UNet++: A nested U-Net architecture for medical image segmentation. In DLMIA (pp. 3-11). https://doi.org/10.1007/978-3-030-00889-5_1",
 "[9] Oktay, O., Schlemper, J., Folgoc, L. L., et al. (2018). Attention U-Net: Learning where to look for the pancreas. arXiv:1804.03999",
 "[10] Huang, H., Lin, L., Tong, R., et al. (2020). UNet 3+: A full-scale connected UNet for medical image segmentation. In ICASSP (pp. 1055-1059). arXiv:2004.08790",
 "[11] Chen, J., Lu, Y., Yu, Q., et al. (2021). TransUNet: Transformers make strong encoders for medical image segmentation. arXiv:2102.04306",
 "[12] Chen, G., Li, L., Dai, Y., Zhang, J., & Yap, M. H. (2023). AAU-Net: An adaptive attention U-Net for breast lesions segmentation in ultrasound images. IEEE Transactions on Medical Imaging, 42(5), 1289-1300. https://doi.org/10.1109/TMI.2022.3226268",
 "[13] Derakhshandeh, P., & Mahloojifar, A. (2024/2025). Modifying the U-Net's encoder-decoder architecture for segmentation of tumors in breast ultrasound images (CResU-Net). Journal of Digital Imaging. arXiv:2409.00647",
 "[14] Latha, M., et al. (2024). Revolutionizing breast ultrasound diagnostics with EfficientNet-B7 and Explainable AI. BMC Medical Imaging, 24, 230. https://doi.org/10.1186/s12880-024-01404-3",
 "[15] Sulaiman, A., et al. (2024). Attention based UNet model for breast cancer segmentation using BUSI dataset. Scientific Reports, 14, 22422. https://doi.org/10.1038/s41598-024-72712-5",
 "[16] Anari, S., et al. (2025). Explainable attention based breast tumor segmentation using a combination of UNet, ResNet, DenseNet, and EfficientNet models. Scientific Reports, 15. https://doi.org/10.1038/s41598-024-84504-y",
 "[17] Bruno, A., et al. (2025). A dual-stage deep learning framework for breast ultrasound image segmentation and classification. Journal of Medical Systems, 49. https://doi.org/10.1007/s10916-025-02298-6",
 "[18] Xian, M., Zhang, Y., Cheng, H. D., et al. (2018). Automatic breast ultrasound image segmentation: A survey. Pattern Recognition, 79, 340-355. https://doi.org/10.1016/j.patcog.2018.02.012",
 "[19] Cheng, H. D., Shan, J., Ju, W., Guo, Y., & Zhang, L. (2010). Automated breast cancer detection and classification using ultrasound images: A survey. Pattern Recognition, 43(1), 299-317. https://doi.org/10.1016/j.patcog.2009.05.012",
 "[20] Isensee, F., Jaeger, P. F., Kohl, S. A. A., Petersen, J., & Maier-Hein, K. H. (2021). nnU-Net: A self-configuring method for deep learning-based biomedical image segmentation. Nature Methods, 18, 203-211. https://doi.org/10.1038/s41592-020-01008-z",
 "[21] Kirillov, A., Mintun, E., Ravi, N., et al. (2023). Segment Anything. In ICCV (pp. 4015-4026). arXiv:2304.02643",
 "[22] Ma, J., He, Y., Li, F., Han, L., You, C., & Wang, B. (2024). Segment anything in medical images (MedSAM). Nature Communications, 15, 654. https://doi.org/10.1038/s41467-024-44824-z",
 "[23] Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). Grad-CAM: Visual explanations from deep networks via gradient-based localization. In ICCV (pp. 618-626). arXiv:1610.02391",
 "[24] Al-Dhabyani, W., Gomaa, M., Khaled, H., & Fahmy, A. (2019). Deep learning approaches for data augmentation and classification of breast masses using ultrasound images. International Journal of Advanced Computer Science and Applications, 10(5). https://doi.org/10.14569/IJACSA.2019.0100579",
 "[25] Lu, Y., et al. (2025). Automatic joint segmentation and classification of breast ultrasound images via multi-task learning with object contextual attention. Frontiers in Oncology, 15, 1567577. https://doi.org/10.3389/fonc.2025.1567577",
]
for rtext in refs:
    p = doc.add_paragraph(rtext)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs: run.font.size = Pt(10)

doc.save(OUT)
# rough word count
wc = sum(len(p.text.split()) for p in doc.paragraphs)
print("saved", OUT)
print("paragraphs:", len(doc.paragraphs), "approx words:", wc)
